from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import DashScopeEmbeddings
from pymilvus import MilvusClient
from src.config import (
    MILVUS_URL,
    EMBEDDING_MODEL,
    DASHSCOPE_API_KEY,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    TOP_K,
    COLLECTION_NAME,
)


class KnowledgeBase:
    """本地知识库，基于 Milvus + DashScope 向量检索。

    懒初始化：仅在首次访问 client 或 embed_model 时才建立连接，
    因此 import 本模块不需要 Milvus 运行。
    """

    def __init__(
        self,
        milvus_url: str,
        collection_name: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        top_k: int = 4,
        embedding_model_name: str = "qwen3.7-text-embedding",
        dashscope_api_key: str | None = None,
    ):
        self.milvus_url = milvus_url
        self.collection_name = collection_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.top_k = top_k
        self.embedding_model_name = embedding_model_name
        self.dashscope_api_key = dashscope_api_key

        self._client = None
        self._embed_model = None
        self._dimension = None

    # ------------------------------------------------------------------
    # 懒初始化属性
    # ------------------------------------------------------------------

    @property
    def client(self) -> MilvusClient:
        if self._client is None:
            self._client = MilvusClient(self.milvus_url)
        return self._client

    @property
    def embed_model(self) -> DashScopeEmbeddings:
        if self._embed_model is None:
            self._embed_model = DashScopeEmbeddings(
                model=self.embedding_model_name,
                dashscope_api_key=self.dashscope_api_key,
            )
        return self._embed_model

    # ------------------------------------------------------------------
    # 内部方法
    # ------------------------------------------------------------------

    def _get_embedding_dim(self) -> int:
        """通过 embed 一条探针字符串来自动检测 embedding 维度，结果缓存。"""
        if self._dimension is None:
            sample = self.embed_model.embed_query("dimension probe")
            self._dimension = len(sample)
        return self._dimension

    def _ensure_collection(self) -> None:
        """确保 Milvus collection 存在（幂等）。"""
        if not self.client.has_collection(self.collection_name):
            dim = self._get_embedding_dim()
            self.client.create_collection(
                collection_name=self.collection_name,
                dimension=dim,
                auto_id=True,
            )

    # ------------------------------------------------------------------
    # 公开方法
    # ------------------------------------------------------------------

    def index_texts(
        self, texts: list[str], sources: list[str] | None = None
    ) -> str:
        """将一批文本切分、embedding 后存入 Milvus。

        Args:
            texts: 待索引的文本列表，每个元素可以是一个完整文档。
            sources: 每个文本的来源标识，长度必须与 texts 相同。
                     未提供时默认用 "unknown"。

        Returns:
            状态字符串，形如 "Indexed 15 chunks from 3 documents into 'kb'."
        """
        if not texts:
            return "No texts provided to index."

        if sources is None:
            sources = ["unknown"] * len(texts)
        elif len(sources) != len(texts):
            raise ValueError(
                f"sources length ({len(sources)}) must match "
                f"texts length ({len(texts)})"
            )

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", "。", ".", " ", ""],
            keep_separator=False,
        )

        all_chunks: list[str] = []
        chunk_sources: list[str] = []
        for text, source in zip(texts, sources):
            chunks = splitter.split_text(text)
            all_chunks.extend(chunks)
            chunk_sources.extend([source] * len(chunks))

        self._ensure_collection()

        vectors = self.embed_model.embed_documents(all_chunks)

        data = [
            {
                "vector": vectors[i],
                "text": all_chunks[i],
                "source": chunk_sources[i],
            }
            for i in range(len(all_chunks))
        ]

        self.client.insert(collection_name=self.collection_name, data=data)

        return (
            f"Indexed {len(all_chunks)} chunks from {len(texts)} documents "
            f"into '{self.collection_name}'."
        )

    def index_file(
        self, file_path: str, source_name: str | None = None
    ) -> str:
        """读取文件并索引到知识库。

        支持格式：.txt, .md, .pdf, .docx

        Args:
            file_path: 文件路径。
            source_name: 来源标识，默认使用文件名（不含扩展名）。

        Returns:
            状态字符串。
        """
        path = Path(file_path)
        if not path.exists():
            return f"File not found: {file_path}"

        if source_name is None:
            source_name = path.stem

        suffix = path.suffix.lower()

        # --- 文本文件 ---
        if suffix in (".txt", ".md", ".markdown", ".rst"):
            text = path.read_text(encoding="utf-8")
            return self.index_texts([text], sources=[source_name])

        # --- PDF ---
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return "pypdf is not installed. Run: pip install pypdf"

            reader = PdfReader(str(path))
            pages: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            if not pages:
                return f"No extractable text found in: {file_path}"
            text = "\n\n".join(pages)
            return self.index_texts([text], sources=[source_name])

        # --- Word ---
        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "python-docx is not installed. Run: pip install python-docx"

            doc = Document(str(path))

            # 段落文本
            parts: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            # 表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())

            if not parts:
                # 最后兜底：尝试直接从 XML body 提取
                body = doc.element.body
                raw_text = "".join(body.itertext()).strip()
                if raw_text:
                    parts.append(raw_text)

            if not parts:
                return (
                    f"No extractable text found in: {file_path}. "
                    "The document may contain only images, scanned pages, or be corrupted."
                )

            text = "\n".join(parts)
            return self.index_texts([text], sources=[source_name])

        return f"Unsupported file format: {suffix}"

    def search(
        self, query: str, top_k: int | None = None
    ) -> list[dict]:
        """在知识库中搜索与 query 最相似的 chunk。

        Args:
            query: 查询文本。
            top_k: 返回条数，默认使用实例的 top_k 配置。

        Returns:
            搜索结果列表，每个元素为 {"id": int, "distance": float,
            "entity": {"text": ..., "source": ...}}。
            无结果时返回空列表。
        """
        if top_k is None:
            top_k = self.top_k

        self._ensure_collection()

        query_vector = self.embed_model.embed_query(query)

        results = self.client.search(
            collection_name=self.collection_name,
            data=[query_vector],
            limit=top_k,
            output_fields=["text", "source"],
        )

        return results[0] if results else []

    def retrieve(
        self, query: str, top_k: int | None = None
    ) -> str:
        """搜索并返回格式化后的检索结果，适合直接展示或传给 Agent。

        Args:
            query: 查询文本。
            top_k: 返回条数。

        Returns:
            格式化字符串，包含编号、来源、相关度和文本片段。
        """
        results = self.search(query, top_k)

        if not results:
            return "No relevant documents found in the knowledge base."

        formatted: list[str] = []
        for i, result in enumerate(results, 1):
            entity = result.get("entity", {})
            text = entity.get("text", "")
            source = entity.get("source", "unknown")
            distance = result.get("distance", 0)
            formatted.append(
                f"{i}. [Source: {source}] (relevance: {distance:.4f})\n"
                f"   {text}"
            )

        return "\n\n".join(formatted)


# ------------------------------------------------------------------
# 模块级便捷函数 —— 使用默认配置快速调用
# ------------------------------------------------------------------

def index_documents(
    texts: list[str],
    sources: list[str] | None = None,
    collection_name: str | None = None,
) -> str:
    """便捷函数：将文档索引到默认知识库。"""
    kb = KnowledgeBase(
        milvus_url=MILVUS_URL,
        collection_name=collection_name or COLLECTION_NAME,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        top_k=TOP_K,
        embedding_model_name=EMBEDDING_MODEL,
        dashscope_api_key=DASHSCOPE_API_KEY,
    )
    return kb.index_texts(texts, sources)


def index_file(
    file_path: str,
    source_name: str | None = None,
    collection_name: str | None = None,
) -> str:
    """便捷函数：将文件索引到默认知识库。

    支持 .txt, .md, .pdf, .docx 格式。
    """
    kb = KnowledgeBase(
        milvus_url=MILVUS_URL,
        collection_name=collection_name or COLLECTION_NAME,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        top_k=TOP_K,
        embedding_model_name=EMBEDDING_MODEL,
        dashscope_api_key=DASHSCOPE_API_KEY,
    )
    return kb.index_file(file_path, source_name)


def retrieve_documents(
    query: str,
    top_k: int | None = None,
    collection_name: str | None = None,
) -> str:
    """便捷函数：在默认知识库中检索文档。"""
    kb = KnowledgeBase(
        milvus_url=MILVUS_URL,
        collection_name=collection_name or COLLECTION_NAME,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        top_k=top_k or TOP_K,
        embedding_model_name=EMBEDDING_MODEL,
        dashscope_api_key=DASHSCOPE_API_KEY,
    )
    return kb.retrieve(query, top_k)
